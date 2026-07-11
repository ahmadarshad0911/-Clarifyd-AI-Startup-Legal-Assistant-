"""Per-user letterhead validation and document compositing.

Two responsibilities:

1. `validate_letterhead` — gate an uploaded file: extension, magic bytes, size,
   and an **A4 page-size check** (the product promises A4 letterheads). Returns
   the classified kind ("pdf" | "docx") or raises ValueError.
2. `render_on_letterhead` — composite generated document text onto the stored
   letterhead. A PDF letterhead is used as the page background with the text
   drawn in a clear body band; a DOCX letterhead is opened directly so its
   header/footer/logo survive and the text is appended to the body.

No external services or new dependencies — pypdf, reportlab, and python-docx are
already in requirements.txt.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from xml.sax.saxutils import escape

# A4 is 210mm x 297mm.
_A4_WIDTH_PT = 595.276
_A4_HEIGHT_PT = 841.89
# Generous tolerance — scanners/exporters round page boxes, and "A4-ish" is
# close enough that rejecting a 594.9pt page would only frustrate the user.
_A4_TOL_PT = 14.0  # ~5mm

_PDF_MAGIC = b"%PDF-"
_DOCX_MAGIC = b"PK\x03\x04"

MAX_LETTERHEAD_BYTES = 5 * 1024 * 1024  # 5MB — header/footer art, not a book.


@dataclass(frozen=True)
class LetterheadMeta:
    kind: str  # "pdf" | "docx"
    mime: str


def _is_a4(width_pt: float, height_pt: float) -> bool:
    """True if the page is A4 in either orientation, within tolerance."""
    portrait = (
        abs(width_pt - _A4_WIDTH_PT) <= _A4_TOL_PT
        and abs(height_pt - _A4_HEIGHT_PT) <= _A4_TOL_PT
    )
    landscape = (
        abs(width_pt - _A4_HEIGHT_PT) <= _A4_TOL_PT
        and abs(height_pt - _A4_WIDTH_PT) <= _A4_TOL_PT
    )
    return portrait or landscape


def _validate_pdf_a4(payload: bytes) -> None:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(payload))
    except Exception as exc:  # corrupt / not a real PDF
        raise ValueError("The letterhead PDF could not be read.") from exc
    if not reader.pages:
        raise ValueError("The letterhead PDF has no pages.")
    box = reader.pages[0].mediabox
    if not _is_a4(float(box.width), float(box.height)):
        raise ValueError(
            "The letterhead must be A4 size (210 × 297 mm). "
            "Please export it as A4 and re-upload."
        )


def _validate_docx_a4(payload: bytes) -> None:
    from docx import Document

    try:
        doc = Document(io.BytesIO(payload))
    except Exception as exc:
        raise ValueError("The letterhead Word file could not be read.") from exc
    if not doc.sections:
        raise ValueError("The letterhead Word file has no page setup.")
    section = doc.sections[0]
    # section.page_width/height are Length values in EMU (1 point = 12700 EMU).
    width_pt = float(section.page_width) / 12700.0 if section.page_width else 0.0
    height_pt = float(section.page_height) / 12700.0 if section.page_height else 0.0
    if not _is_a4(width_pt, height_pt):
        raise ValueError(
            "The letterhead must be A4 size (210 × 297 mm). "
            "Please set the page size to A4 and re-upload."
        )


def validate_letterhead(file_name: str, payload: bytes) -> LetterheadMeta:
    """Validate an uploaded letterhead. Raise ValueError with a user-facing
    message on any problem, else return its classified kind + mime."""
    if not payload:
        raise ValueError("The letterhead file is empty.")
    if len(payload) > MAX_LETTERHEAD_BYTES:
        mb = MAX_LETTERHEAD_BYTES // (1024 * 1024)
        raise ValueError(f"The letterhead must be under {mb}MB.")

    name = (file_name or "").lower().strip()
    if payload.startswith(_PDF_MAGIC) and name.endswith(".pdf"):
        _validate_pdf_a4(payload)
        return LetterheadMeta(kind="pdf", mime="application/pdf")
    if payload.startswith(_DOCX_MAGIC) and name.endswith(".docx"):
        _validate_docx_a4(payload)
        return LetterheadMeta(
            kind="docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    raise ValueError("The letterhead must be a PDF or Word (.docx) file.")


def _split_paragraphs(text: str) -> list[str]:
    """Blank-line separated blocks become paragraphs; single newlines inside a
    block are preserved as line breaks."""
    blocks: list[str] = []
    for block in text.replace("\r\n", "\n").split("\n\n"):
        stripped = block.strip()
        if stripped:
            blocks.append(stripped)
    return blocks


def _render_pdf_on_pdf(letterhead: bytes, title: str, body: str) -> bytes:
    from pypdf import PdfReader, PdfWriter
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    # Body band: clear the top (letterhead header/logo) and bottom (footer).
    text_buf = io.BytesIO()
    doc = SimpleDocTemplate(
        text_buf,
        pagesize=A4,
        topMargin=2.2 * inch,
        bottomMargin=1.5 * inch,
        leftMargin=1.0 * inch,
        rightMargin=1.0 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"], fontSize=16, spaceAfter=18, alignment=TA_LEFT
    )
    body_style = ParagraphStyle(
        "DocBody", parent=styles["BodyText"], fontSize=10.5, leading=16, spaceAfter=10
    )
    story: list = []
    if title.strip():
        story.append(Paragraph(escape(title.strip()), title_style))
    for para in _split_paragraphs(body):
        story.append(Paragraph(escape(para).replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 2))
    doc.build(story)
    text_buf.seek(0)

    # Underlay: stamp each rendered body page over a fresh copy of the
    # letterhead page so the header/footer art shows through on every page.
    body_reader = PdfReader(text_buf)
    writer = PdfWriter()
    for body_page in body_reader.pages:
        base = PdfReader(io.BytesIO(letterhead)).pages[0]
        base.merge_page(body_page)
        writer.add_page(base)

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def _render_docx_on_docx(letterhead: bytes, title: str, body: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(io.BytesIO(letterhead))
    if title.strip():
        heading = doc.add_paragraph()
        run = heading.add_run(title.strip())
        run.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for para in _split_paragraphs(body):
        p = doc.add_paragraph()
        # Preserve single newlines within a block as line breaks.
        lines = para.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render_on_letterhead(kind: str, letterhead: bytes, title: str, body: str) -> tuple[bytes, str, str]:
    """Composite `body` (the generated document text) onto `letterhead`.

    Returns (file_bytes, mime, extension). PDF letterhead -> PDF; DOCX -> DOCX.
    """
    if kind == "pdf":
        return _render_pdf_on_pdf(letterhead, title, body), "application/pdf", "pdf"
    if kind == "docx":
        data = _render_docx_on_docx(letterhead, title, body)
        return (
            data,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    raise ValueError(f"Unsupported letterhead kind: {kind!r}")
